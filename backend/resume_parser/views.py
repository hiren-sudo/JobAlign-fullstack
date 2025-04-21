# from django.shortcuts import render
# from .models import Resume
# import re
# import pdfminer.high_level
# import docx
# import pytesseract
# from pdf2image import convert_from_bytes
# from PIL import Image
# import io

# # Path to Tesseract executable
# pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# # PDF text extraction
# def extract_text_from_pdf(pdf_file):
#     try:
#         pdf_bytes = pdf_file.read()
#         pdf_stream = io.BytesIO(pdf_bytes)
#         return pdfminer.high_level.extract_text(pdf_stream)
#     except Exception as e:
#         print(f"Error extracting text from PDF: {e}")
#         return ""

# # OCR for scanned PDFs
# def extract_text_with_ocr(pdf_file):
#     try:
#         pdf_bytes = pdf_file.read()
#         images = convert_from_bytes(pdf_bytes)
#         text = ""
#         for img in images:
#             text += pytesseract.image_to_string(img, config='--psm 6')
#         return text
#     except Exception as e:
#         print(f"OCR extraction failed: {e}")
#         return ""

# # DOCX text extraction
# def extract_text_from_docx(docx_file):
#     try:
#         doc = docx.Document(docx_file)
#         return "\n".join([para.text for para in doc.paragraphs])
#     except Exception as e:
#         print(f"Error reading DOCX: {e}")
#         return ""

# # Clean and preprocess text
# def preprocess_text(text):
#     text = text.replace('\r', '\n').replace('\t', ' ').strip()
#     text = re.sub(r'\s+', ' ', text)
#     text = re.sub(r'[^a-zA-Z0-9@.,+()&:;\s-]', '', text)
#     return text

# # Regex-based field extraction
# def extract_name(text):
#     name_match = re.search(r'\b(Name[:\s]*)?([A-Z][a-z]+(?:\s[A-Z][a-z]+)*)', text)
#     return name_match.group(2) if name_match else "N/A"

# def extract_email(text):
#     match = re.search(r'[\w\.-]+@[\w\.-]+', text)
#     return match.group(0) if match else "N/A"

# def extract_phone(text):
#     match = re.search(r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', text)
#     return match.group(0) if match else "N/A"

# def extract_skills(text):
#     skills_keywords = ["Python", "Django", "React", "JavaScript", "SQL", "AI", "Machine Learning"]
#     skills = [skill for skill in skills_keywords if re.search(r'\b' + re.escape(skill) + r'\b', text, re.IGNORECASE)]
#     return ', '.join(skills) if skills else "N/A"

# def extract_education(text):
#     degrees = ["B.Tech", "Bachelor", "M.Tech", "MSc", "PhD", "Master", "BSc", "Engineering", "Computer Science"]
#     education = [degree for degree in degrees if re.search(r'\b' + re.escape(degree) + r'\b', text, re.IGNORECASE)]
#     return ', '.join(education) if education else "N/A"

# def extract_experience(text):
#     match = re.findall(r'(\d+)\s*(?:years?|months?)', text, re.IGNORECASE)
#     return ', '.join(f"{m} years" for m in match) if match else "N/A"

# # Main parsing logic
# def upload_resume(request):
#     if request.method == 'POST':
#         uploaded_file = request.FILES['resume']
#         file_type = uploaded_file.name.split('.')[-1].lower()

#         if file_type == 'pdf':
#             file_text = extract_text_from_pdf(uploaded_file)
#             if not file_text.strip():
#                 uploaded_file.seek(0)
#                 file_text = extract_text_with_ocr(uploaded_file)
#         elif file_type == 'docx':
#             file_text = extract_text_from_docx(uploaded_file)
#         else:
#             file_text = "Unsupported file format."

#         print("Extracted text:", file_text)

#         cleaned_text = preprocess_text(file_text)

#         parsed_data = {
#             'skills': extract_skills(cleaned_text),
#             'education': extract_education(cleaned_text),
#             'experience': extract_experience(cleaned_text)
#         }

#         Resume.objects.create(**parsed_data)

#         return render(request, 'resume_parser/resume_parsed.html', {'parsed_data': parsed_data})

#     return render(request, 'resume_parser/upload_resume.html')






from django.shortcuts import render, redirect
from django.core.exceptions import ValidationError
from django.contrib import messages
from .models import Resume
import re
import pdfminer.high_level
import docx
import pytesseract
from pdf2image import convert_from_path
from PIL import Image
import io
import json
import os
import google.generativeai as genai
from django.conf import settings
from django.core.files.uploadedfile import UploadedFile
import logging
import tempfile

# Set up logging
logger = logging.getLogger(__name__)

# Configure Gemini API
try:
    genai.configure(api_key=settings.GOOGLE_GEMINI_API_KEY)
except AttributeError:
    logger.error("Google Gemini API key is missing in settings.")
    raise Exception("Google Gemini API key not configured.")

# Cache file for skills and degrees
CACHE_FILE = os.path.join(settings.BASE_DIR, "cached_skills_degrees.json")
ALLOWED_EXTENSIONS = {'pdf', 'docx'}
MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB

# # Load cached skills and degrees
# def get_cached_data():
#     """Load skills and degrees from cache or fetch from Gemini API if not available."""
#     if os.path.exists(CACHE_FILE):
#         try:
#             with open(CACHE_FILE, "r") as file:
#                 return json.load(file)
#         except (json.JSONDecodeError, IOError) as e:
#             logger.error(f"Failed to load cached data: {e}")
#             return fetch_skills_degrees_from_gemini()
#     return fetch_skills_degrees_from_gemini()

# def fetch_skills_degrees_from_gemini():
#     """Fetch skills and degrees from Gemini API and cache them."""
#     prompt = """
#     Generate two JSON lists:
#     - "skills": A list of 5000+ job-related skills across all industries.
#     - "degrees": A list of 2000+ academic degrees worldwide.
#     Return only valid JSON format.
#     """
#     try:
#         model = genai.GenerativeModel("gemini-1.5-pro-latest")
#         response = model.generate_content(prompt)
#         # Clean the response to extract JSON
#         cleaned_response = response.text.strip()
#         cleaned_response = re.sub(r'^```json\s*|\s*```$', '', cleaned_response, flags=re.MULTILINE)
#         cleaned_response = cleaned_response.strip()
#         # Parse the cleaned response as JSON
#         data = json.loads(cleaned_response)
#         if not isinstance(data, dict) or "skills" not in data or "degrees" not in data:
#             logger.error(f"Gemini API returned invalid data: {data}")
#             return {"skills": [], "degrees": []}
#         logger.debug(f"Writing to cache file: {CACHE_FILE}")
#         with open(CACHE_FILE, "w") as file:
#             json.dump(data, file, indent=4)
#         logger.info(f"Successfully created cache file: {CACHE_FILE}")
#         return data
#     except json.JSONDecodeError as e:
#         logger.error(f"JSON parsing error: {e}, Raw response: {response.text}")
#         return {"skills": [], "degrees": []}
#     except Exception as e:
#         logger.error(f"Gemini API error: {e}")
#         return {"skills": [], "degrees": []}

# CACHED_DATA = get_cached_data()
# ALL_SKILLS = set(CACHED_DATA.get("skills", []))
# ALL_DEGREES = set(CACHED_DATA.get("degrees", []))


def get_cached_data():
    """Load skills and degrees from cache or fetch from Gemini API if not available."""
    if os.path.exists(CACHE_FILE):
        try:
            with open(CACHE_FILE, "r") as file:
                logger.debug(f"Loading cached data from {CACHE_FILE}")
                data = json.load(file)
                logger.debug(f"Loaded data: {data}")
                return data
        except (json.JSONDecodeError, IOError) as e:
            logger.error(f"Failed to load cached data: {e}")
            return fetch_skills_degrees_from_gemini()
    logger.info(f"Cache file {CACHE_FILE} not found. Fetching from Gemini API.")
    return fetch_skills_degrees_from_gemini()

def fetch_skills_degrees_from_gemini():
    """Fetch skills and degrees from Gemini API and cache them."""
    prompt = """
    Generate two JSON lists and return the result as a valid JSON string only, with no additional text, markdown, or formatting:
    - "skills": A list of 5000 job-related skills across all industries, including both soft skills (e.g., Communication, Leadership) and technical skills (e.g., Python, Java, HTML, CSS, JavaScript, SQL, Django, React, Git, Github, MySQL, MongoDB, Tailwind CSS, Swift, Go, R, Node.js, AWS, Docker, Kubernetes, Machine Learning, Data Science).
    - "degrees": A list of 150 academic degrees worldwide, including variations of engineering degrees (e.g., Bachelor of Engineering, BEng, Master of Engineering, MEng, Computer Engineering) and other common degrees (e.g., B.Tech, M.Sc, PhD, B.A., M.A.).
    """
    try:
        logger.debug("Attempting to fetch data from Gemini API...")
        model = genai.GenerativeModel("gemini-1.5-pro-latest")
        response = model.generate_content(prompt)
        logger.debug(f"Raw response from Gemini API: {response.text}")
        # Clean the response to extract JSON
        cleaned_response = response.text.strip()
        cleaned_response = re.sub(r'```json\s*\n|\n\s*```', '', cleaned_response, flags=re.MULTILINE)
        cleaned_response = re.sub(r'```\s*\n|\n\s*```', '', cleaned_response, flags=re.MULTILINE)
        cleaned_response = re.sub(r'^.*?{', '{', cleaned_response, flags=re.DOTALL)
        cleaned_response = re.sub(r'}[^}]*$', '}', cleaned_response, flags=re.DOTALL)
        cleaned_response = cleaned_response.strip()
        logger.debug(f"Cleaned response: {cleaned_response}")
        # Parse the cleaned response as JSON
        data = json.loads(cleaned_response)
        logger.debug(f"Parsed JSON data: {data}")
        if not isinstance(data, dict) or "skills" not in data or "degrees" not in data:
            logger.error(f"Gemini API returned invalid data: {data}")
            return {"skills": [], "degrees": []}
        logger.debug(f"Writing to cache file: {CACHE_FILE}")
        with open(CACHE_FILE, "w") as file:
            json.dump(data, file, indent=4)
        logger.info(f"Successfully created cache file: {CACHE_FILE}")
        return data
    except json.JSONDecodeError as e:
        logger.error(f"JSON parsing error: {e}, Raw response: {response.text}")
        return {"skills": [], "degrees": []}
    except Exception as e:
        logger.error(f"Gemini API error: {e}")
        return {"skills": [], "degrees": []}

CACHED_DATA = get_cached_data()
ALL_SKILLS = set(CACHED_DATA.get("skills", []))
ALL_DEGREES = set(CACHED_DATA.get("degrees", []))
logger.debug(f"ALL_SKILLS: {ALL_SKILLS}")
logger.debug(f"ALL_DEGREES: {ALL_DEGREES}")


# Path to Tesseract executable (make configurable)
TESSERACT_CMD = getattr(settings, 'TESSERACT_CMD', r'C:\Program Files\Tesseract-OCR\tesseract.exe')
pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD

# File validation
def validate_file(uploaded_file: UploadedFile):
    """Validate file type and size."""
    if not uploaded_file:
        raise ValidationError("No file uploaded.")
    
    file_type = uploaded_file.name.split('.')[-1].lower()
    if file_type not in ALLOWED_EXTENSIONS:
        raise ValidationError(f"Unsupported file format: {file_type}. Only PDF and DOCX are allowed.")
    
    if uploaded_file.size > MAX_FILE_SIZE:
        raise ValidationError(f"File too large. Maximum size allowed is {MAX_FILE_SIZE / (1024 * 1024)} MB.")

# PDF text extraction
def extract_text_from_pdf(pdf_file: UploadedFile) -> str:
    """Extract text from a PDF file using pdfminer."""
    try:
        pdf_bytes = pdf_file.read()
        pdf_stream = io.BytesIO(pdf_bytes)
        text = pdfminer.high_level.extract_text(pdf_stream)
        return text if text else ""
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        return ""

# OCR for scanned PDFs
def extract_text_with_ocr(pdf_file: UploadedFile) -> str:
    """Extract text from a PDF using OCR if direct extraction fails."""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
            for chunk in pdf_file.chunks():
                temp_file.write(chunk)
            temp_file_path = temp_file.name

        images = convert_from_path(temp_file_path)
        text = ""
        for img in images:
            text += pytesseract.image_to_string(img, config='--psm 6')
        os.remove(temp_file_path)  # Clean up
        return text if text else ""
    except Exception as e:
        logger.error(f"OCR extraction failed: {e}")
        return ""
    finally:
        if 'temp_file_path' in locals() and os.path.exists(temp_file_path):
            os.remove(temp_file_path)

# DOCX text extraction
def extract_text_from_docx(docx_file: UploadedFile) -> str:
    """Extract text from a DOCX file."""
    try:
        doc = docx.Document(docx_file)
        return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    except Exception as e:
        logger.error(f"Error reading DOCX: {e}")
        return ""

# Clean and preprocess text
def preprocess_text(text: str) -> str:
    """Clean and preprocess extracted text."""
    if not text:
        return ""
    text = text.replace('\r', '\n').replace('\t', ' ').strip()
    text = re.sub(r'\s+', ' ', text)  # Normalize whitespace
    text = re.sub(r'[^a-zA-Z0-9@.,+()&:;\s-]', '', text)  # Remove special characters
    return text

# Field extraction
def extract_skills(text: str) -> str:
    """Extract skills by matching against the cached skills list."""
    if not ALL_SKILLS:
        logger.warning("Skills list is empty. Check Gemini API or cache.")
        return "N/A"
    text_lower = text.lower()
    matched_skills = set()
    for skill in ALL_SKILLS:
        skill_lower = skill.lower()
        # Match whole words to avoid partial matches (e.g., "Java" in "JavaScript")
        # Use word boundaries and ensure the skill is a standalone word or phrase
        pattern = r'\b' + re.escape(skill_lower) + r'\b'
        if re.search(pattern, text_lower):
            matched_skills.add(skill)
    logger.debug(f"Matched Skills: {matched_skills}")
    return ', '.join(sorted(matched_skills)) if matched_skills else "N/A"

def extract_education(text: str) -> str:
    """Extract education by matching against the cached degrees list."""
    if not ALL_DEGREES:
        logger.warning("Degrees list is empty. Check Gemini API or cache.")
        return "N/A"
    text_lower = text.lower()
    matched_degrees = set()
    for degree in ALL_DEGREES:
        degree_lower = degree.lower()
        # Match whole phrases with word boundaries
        pattern = r'\b' + re.escape(degree_lower) + r'\b'
        if re.search(pattern, text_lower):
            matched_degrees.add(degree)
        # Handle variations for engineering degrees
        elif "bachelor of engineering" in degree_lower and "bachelor of engineering" in text_lower:
            matched_degrees.add(degree)
        elif "beng" in degree_lower and "bachelor of engineering" in text_lower:
            matched_degrees.add(degree)
        # Match "Computer Engineering" specifically
        elif "computer engineering" in degree_lower and "computer engineering" in text_lower:
            matched_degrees.add(degree)
    logger.debug(f"Matched Degrees: {matched_degrees}")
    return ', '.join(sorted(matched_degrees)) if matched_degrees else "N/A"

# Main parsing logic
def upload_resume(request):
    """Handle resume upload, parsing, and storage."""
    if request.method != 'POST':
        try:
            return render(request, 'resume_parser/upload_resume.html')
        except Exception as e:
            logger.error(f"Error rendering upload_resume.html: {e}")
            messages.error(request, "Error loading upload page. Please try again.")
            return redirect('upload_resume')

    try:
        # Validate file
        uploaded_file = request.FILES.get('resume')
        validate_file(uploaded_file)

        file_type = uploaded_file.name.split('.')[-1].lower()
        file_text = ""

        # Extract text based on file type
        if file_type == 'pdf':
            file_text = extract_text_from_pdf(uploaded_file)
            if not file_text.strip():
                uploaded_file.seek(0)  # Reset file pointer
                file_text = extract_text_with_ocr(uploaded_file)
        elif file_type == 'docx':
            file_text = extract_text_from_docx(uploaded_file)

        if not file_text.strip():
            messages.error(request, "Unable to extract text from the file. Please try a different file.")
            return redirect('upload_resume')

        # Preprocess and parse the text
        cleaned_text = preprocess_text(file_text)
        logger.debug(f"Cleaned Text: {cleaned_text}")

        parsed_data = {
            'skills': extract_skills(cleaned_text) or "N/A",
            'education': extract_education(cleaned_text) or "N/A",
        }

        logger.debug(f"Parsed Data: {parsed_data}")

        # Save to database with validation
        try:
            resume_instance = Resume(**parsed_data)
            resume_instance.full_clean()  # Validate the instance
            resume_instance.save()
        except ValidationError as ve:
            logger.error(f"Validation error saving resume: {ve}")
            messages.error(request, f"Invalid data in resume: {str(ve)}. Please check the file and try again.")
            return redirect('upload_resume')
        except Exception as e:
            logger.error(f"Error saving to database: {e}")
            messages.error(request, "Error saving resume data. Please try again.")
            return redirect('upload_resume')

        messages.success(request, "Resume parsed and saved successfully!")
        try:
            return render(request, 'resume_parser/resume_parsed.html', {'parsed_data': parsed_data})
        except Exception as e:
            logger.error(f"Error rendering resume_parsed.html: {e}")
            messages.error(request, "Resume saved, but there was an error displaying the results.")
            return redirect('upload_resume')

    except ValidationError as ve:
        messages.error(request, str(ve))
        return redirect('upload_resume')
    except Exception as e:
        logger.error(f"Unexpected error in upload_resume: {e}")
        messages.error(request, "An unexpected error occurred. Please try again later.")
        return redirect('upload_resume')